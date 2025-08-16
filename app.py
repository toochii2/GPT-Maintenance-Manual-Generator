import streamlit as st
import tempfile
import os
from pathlib import Path
import time
from openai import OpenAI
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import io
from pydub import AudioSegment
import moviepy.editor as mp
import tempfile
import cv2
import numpy as np
from PIL import Image
import base64
from datetime import timedelta
import json
from difflib import SequenceMatcher
import shutil  # Add this import at the top with other imports

# Set page configuration with expanded layout by default
st.set_page_config(
    page_title="Maintenance Manual Generator",
    page_icon="🔧",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Helper timestamp functions
def normalize_timestamp(timestamp):
    """
    Normalize timestamp format to ensure consistent matching.
    Handles formats like '1:23' vs '01:23' and removes any leading zeros.
    """
    # Split by colon
    parts = timestamp.split(':')
    
    # Convert each part to integer to remove leading zeros
    normalized_parts = [str(int(part)) for part in parts]
    
    # Join back with colons
    return ':'.join(normalized_parts)

def timestamp_to_seconds(timestamp):
    """Convert a timestamp to seconds for comparison"""
    parts = timestamp.split(':')
    if len(parts) == 2:  # MM:SS
        return int(parts[0]) * 60 + int(parts[1])
    elif len(parts) == 3:  # HH:MM:SS
        return int(parts[0]) * 3600 + int(parts[1]) * 60 + int(parts[2])
    return 0

#############
# SIDEBAR UI #
#############

st.sidebar.image("logo/Maintenance Manual Generator.png", use_container_width=True)
st.sidebar.title("Settings & Controls")

# Create sidebar sections with expandable areas
api_section = st.sidebar.expander("📋 API Settings", expanded=True)
with api_section:
    api_key = st.text_input("OpenAI API Key:", type="password", help="Required for transcription and AI processing")
    if api_key:
        st.session_state["openai_api_key"] = api_key
        st.success("✅ API Key set")

# File upload section
file_section = st.sidebar.expander("📁 File Input", expanded=True)
with file_section:
    uploaded_file = st.file_uploader("Upload MP3 or MP4 file", type=["mp3", "mp4"])
    
    if uploaded_file is not None:
        st.success(f"✅ Uploaded: {uploaded_file.name}")
        
        # Display file details in a more compact format
        col1, col2 = st.columns(2)
        with col1:
            st.metric("File size", f"{uploaded_file.size / 1024 / 1024:.1f} MB")
        with col2:
            st.metric("File type", uploaded_file.name.split('.')[-1].upper())

# Model controls section
model_section = st.sidebar.expander("🧠 AI Model Controls", expanded=False)
with model_section:
    # LLM Model selection
    llm_model = st.selectbox(
        "LLM Model:",
        options=["gpt-4o-mini", "gpt-4o"],
        index=0,
        help="Select which OpenAI model to use. More advanced models may provide better results but cost more credits."
    )
    
    # Temperature control
    temperature = st.slider(
        "Temperature:", 
        min_value=0.0, 
        max_value=1.0, 
        value=0.3,
        step=0.1,
        help="Lower values (0.0-0.3) produce more consistent, focused output. Higher values allow more creativity but may introduce inaccuracies."
    )
    
    # Fact-checking strictness
    fact_checking = st.radio(
        "Fact-checking level:",
        options=["Strict (only transcript facts)", "Balanced (minimal inference)", "Flexible (allow helpful additions)"],
        index=1,
        help="Controls how strictly the system adheres to the original transcript content."
    )
    
    st.write("---")
    
    # Advanced model settings
    col1, col2 = st.columns(2)
    with col1:
        enable_verification = st.checkbox(
            "Content verification",
            value=True,
            help="Adds an additional verification step to reduce hallucinations"
        )
    with col2:
        highlight_inferences = st.checkbox(
            "Highlight inferences",
            value=True,
            help="Marks content that goes beyond the transcript"
        )

# Output options section
output_section = st.sidebar.expander("📄 Output Options", expanded=False)
with output_section:
    # Document formatting options
    manual_title = st.text_input(
        "Manual title:",
        value="",
        placeholder="Enter a title (optional)",
        help="Leave blank to use default title based on filename"
    )
    
    # Image options (only shown for videos)
    if uploaded_file is not None and uploaded_file.name.endswith('.mp4'):
        include_images = st.checkbox("Include images in document", value=True)
        image_size = st.slider("Image size (inches):", 2.0, 6.0, 4.0, 0.5)
        
        # Advanced image matching
        enable_advanced_matching = st.checkbox(
            "Advanced timestamp matching", 
            value=True,
            help="Enable more aggressive image matching techniques"
        )
        
        # Timestamp format standardization
        st.radio(
            "Timestamp format:",
            options=["Auto-detect", "MM:SS", "HH:MM:SS"],
            index=0,
            key="timestamp_format",
            help="Specify the format used in your video"
        )
    else:
        include_images = False
        image_size = 4.0
        enable_advanced_matching = True
    
    # Timestamp position option
    timestamp_position = st.radio(
        "Timestamp position:",
        options=["End of steps", "Beginning of steps"],
        index=0,
        help="Where to place timestamps in relation to instructions"
    )
    
    # Include original language transcript option
    include_original_transcript = st.checkbox(
        "Include original language transcript",
        value=False,
        help="Include both original and translated transcripts in the document (only applies when translation occurs)"
    )

# Processing options section
process_section = st.sidebar.expander("⚙️ Processing Options", expanded=False)
with process_section:
    # Show intermediate results
    show_intermediate = st.checkbox("Show intermediate results", value=True)
    
    # Auto-start processing
    if uploaded_file is not None:
        auto_process = st.checkbox("Auto-start processing", value=False)

# Help & About section
help_section = st.sidebar.expander("❓ Help & About", expanded=False)
with help_section:
    st.markdown("""
    ### How to use this app
    1. Enter your OpenAI API key
    2. Upload an MP3 or MP4 file
    3. Adjust settings as needed
    4. Process the file to generate a manual
    5. Download the resulting Word document
    
    ### About
    This app uses AI to convert maintenance and repair videos into structured instruction manuals with timestamps and images.
    
    Version 1.0.0 | © 2025
    """)

# Initialize OpenAI client
def initialize_client():
    api_key = st.session_state.get("openai_api_key", "")
    if (api_key):
        return OpenAI(api_key=api_key)
    return None

# Get client
client = initialize_client()

# Main content area
st.title("🔧 Maintenance Manual Generator")

# Initialize session state for tracking progress
if 'processing_stage' not in st.session_state:
    st.session_state.processing_stage = None

# Progress tracking
def update_progress(stage, status="running"):
    st.session_state.processing_stage = stage
    st.session_state[f"{stage}_status"] = status

# Check if we should display the process button
if uploaded_file is not None and client:
    if auto_process and st.session_state.processing_stage is None:
        start_processing = True
    elif st.session_state.processing_stage is None:
        start_processing = st.button("▶️ Start Processing", use_container_width=True, type="primary")
    else:
        start_processing = False
else:
    start_processing = False

# Diagnostic functions
def debug_timestamp_formats(timestamps, frames):
    """Show detailed debug information about timestamp formats for troubleshooting"""
    
    if not show_intermediate:
        return
    
    st.subheader("📊 Timestamp Diagnostics")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.write("📹 **Video Frame Timestamps**")
        if frames:
            # Get sample of timestamp formats
            sample_video_ts = list(frames.keys())[:5]
            st.code("\n".join(sample_video_ts))
            
            # Show first digits
            if sample_video_ts:
                first_chars = [ts[0] for ts in sample_video_ts]
                st.write(f"First digits: `{'`.`'.join(first_chars)}`")
                
                # Check if some have leading zeros and others don't
                has_leading_zero = any(ts.startswith('0') for ts in sample_video_ts)
                no_leading_zero = any(not ts.startswith('0') and ts[0].isdigit() for ts in sample_video_ts)
                
                if has_leading_zero and no_leading_zero:
                    st.warning("⚠️ Inconsistent leading zeros detected in video timestamps")
    
    with col2:
        st.write("📝 **Document Timestamps**")
        if timestamps:
            # Get sample of document timestamps
            doc_ts = [ts["start"] for ts in timestamps[:5]]
            st.code("\n".join(doc_ts))
            
            # Show first digits
            if doc_ts:
                first_chars = [ts[0] for ts in doc_ts]
                st.write(f"First digits: `{'`.`'.join(first_chars)}`")
                
                # Check if some have leading zeros and others don't
                has_leading_zero = any(ts.startswith('0') for ts in doc_ts)
                no_leading_zero = any(not ts.startswith('0') and ts[0].isdigit() for ts in doc_ts)
                
                if has_leading_zero and no_leading_zero:
                    st.warning("⚠️ Inconsistent leading zeros detected in document timestamps")
    
    # Check for format compatibility
    if frames and timestamps:
        video_ts_sample = list(frames.keys())[:1]
        doc_ts_sample = [ts["start"] for ts in timestamps[:1]]
        
        if video_ts_sample and doc_ts_sample:
            video_format = video_ts_sample[0].count(':')
            doc_format = doc_ts_sample[0].count(':')
            
            if video_format != doc_format:
                st.error(f"⚠️ Format mismatch: Video uses {'HH:MM:SS' if video_format == 2 else 'MM:SS'} format but document uses {'HH:MM:SS' if doc_format == 2 else 'MM:SS'} format")
            
            # Show normalized versions
            st.write("**Normalized sample comparison:**")
            if video_ts_sample and doc_ts_sample:
                st.code(f"Video: {video_ts_sample[0]} → {normalize_timestamp(video_ts_sample[0])}\nDoc: {doc_ts_sample[0]} → {normalize_timestamp(doc_ts_sample[0])}")

def visualize_timestamp_matches(document_timestamps, available_frames):
    """Visualize which document timestamps match with available frames"""
    if not show_intermediate:
        return
    
    st.subheader("📊 Timestamp Matching Visualization")
    
    # Extract timestamps from the document
    doc_ts = []
    for ts in document_timestamps:
        doc_ts.append(ts["start"])
    
    # Get frame timestamps
    frame_ts = list(available_frames.keys())
    
    # Create normalized versions for both
    norm_doc_ts = [normalize_timestamp(ts) for ts in doc_ts]
    norm_frame_ts = [normalize_timestamp(ts) for ts in frame_ts]
    
    # Find matches and mismatches
    direct_matches = []
    normalized_matches = []
    unmatched = []
    
    for ts in doc_ts:
        if ts in frame_ts:
            direct_matches.append(ts)
        elif normalize_timestamp(ts) in norm_frame_ts:
            normalized_matches.append(ts)
        else:
            unmatched.append(ts)
    
    # Display results
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("Direct matches", len(direct_matches))
        if direct_matches:
            st.write("Sample direct matches:")
            st.code("\n".join(direct_matches[:5]))
    
    with col2:
        st.metric("Normalized matches", len(normalized_matches))
        if normalized_matches:
            st.write("Sample normalized matches:")
            for ts in normalized_matches[:5]:
                st.code(f"{ts} → {normalize_timestamp(ts)}")
    
    with col3:
        st.metric("Unmatched timestamps", len(unmatched))
        if unmatched:
            st.write("Sample unmatched:")
            st.code("\n".join(unmatched[:5]))
    
    # Visualization
    total_docs = len(doc_ts)
    matched_percent = ((len(direct_matches) + len(normalized_matches)) / total_docs) * 100 if total_docs > 0 else 0
    
    st.progress(matched_percent / 100)
    st.write(f"Match rate: {matched_percent:.1f}% ({len(direct_matches) + len(normalized_matches)} of {total_docs})")

# Function to extract audio from video if needed
def extract_audio(video_path):
    update_progress("extract_audio")
    with st.spinner("Extracting audio from video..."):
        temp_audio_file = tempfile.NamedTemporaryFile(delete=False, suffix='.mp3')
        video = mp.VideoFileClip(video_path)
        video.audio.write_audiofile(temp_audio_file.name, verbose=False, logger=None)
        update_progress("extract_audio", "complete")
        return temp_audio_file.name

# Function to extract frames from video at specific timestamps
def extract_frames_at_timestamps(video_path, timestamps):
    update_progress("extract_frames")
    with st.spinner("Extracting images from video at key timestamps..."):
        frames = {}
        normalized_frames = {}  # For easier lookup
        temp_dir = tempfile.mkdtemp()
        
        # Load video
        video = cv2.VideoCapture(video_path)
        fps = video.get(cv2.CAP_PROP_FPS)
        total_frames = int(video.get(cv2.CAP_PROP_FRAME_COUNT))
        duration = total_frames / fps if fps > 0 else 0
        
        # Debug info
        if show_intermediate:
            st.info(f"Video FPS: {fps}, Total frames: {total_frames}, Duration: {duration:.2f} seconds")
            st.info(f"Extracting frames for {len(timestamps)} timestamps")
        
        # Extract frames at regular intervals as a fallback
        interval_frames = {}
        interval = 5  # seconds
        for second in range(0, int(duration), interval):
            frame_number = int(second * fps)
            if frame_number < total_frames:
                video.set(cv2.CAP_PROP_POS_FRAMES, frame_number)
                success, frame = video.read()
                if success:
                    # Format timestamp as MM:SS
                    mins = second // 60
                    secs = second % 60
                    ts = f"{mins}:{secs:02d}"
                    img_path = os.path.join(temp_dir, f"interval_frame_{mins}_{secs}.jpg")
                    cv2.imwrite(img_path, frame)
                    interval_frames[ts] = img_path
        
        # Extract frames from specific timestamps
        extracted_count = 0
        for ts in timestamps:
            # Convert timestamp to seconds
            parts = ts["start"].split(':')
            seconds = 0
            if len(parts) == 2:
                # Format MM:SS
                seconds = int(parts[0]) * 60 + int(parts[1])
            elif len(parts) == 3:
                # Format HH:MM:SS
                seconds = int(parts[0]) * 3600 + int(parts[1]) * 60 + int(parts[2])
            else:
                continue
                
            # Calculate frame number
            frame_number = int(seconds * fps)
            
            # Ensure frame number is valid
            if frame_number < total_frames:
                # Set video to specific frame
                video.set(cv2.CAP_PROP_POS_FRAMES, frame_number)
                success, frame = video.read()
                
                if success:
                    # Save frame as image
                    img_path = os.path.join(temp_dir, f"frame_{ts['start'].replace(':', '_')}.jpg")
                    cv2.imwrite(img_path, frame)
                    
                    # Store with both original and normalized format
                    frames[ts["start"]] = img_path
                    
                    # Also store with normalized key for easier matching
                    norm_ts = normalize_timestamp(ts["start"])
                    normalized_frames[norm_ts] = img_path
                    
                    extracted_count += 1
        
        video.release()
        
        # If no frames were extracted from timestamps, use the interval frames
        if extracted_count == 0 and interval_frames:
            frames = interval_frames
            st.warning(f"No frames extracted from timestamps. Using {len(interval_frames)} frames at {interval}-second intervals instead.")
        else:
            if show_intermediate:
                st.success(f"Successfully extracted {extracted_count} images from timestamps")
        
        # Debug: show what timestamps we have images for
        if frames and show_intermediate:
            st.info(f"Extracted frames for timestamps: {', '.join(frames.keys())}")
        
        update_progress("extract_frames", "complete")
        
        # Return both regular and normalized frames
        return frames, normalized_frames

# Function to display images in Streamlit
def display_extracted_images(frames):
    if not frames:
        return
    
    st.subheader("Extracted Images from Key Timestamps")
    
    # Calculate number of columns (3 images per row)
    num_images = len(frames)
    cols = st.columns(min(3, num_images))
    
    # Display images in a grid layout
    for i, (timestamp, img_path) in enumerate(frames.items()):
        col_idx = i % 3
        with cols[col_idx]:
            img = Image.open(img_path)
            st.image(img, caption=f"Timestamp: {timestamp}", use_container_width=True)

# Function to transcribe audio using OpenAI's Whisper API
def transcribe_audio(client, audio_path):
    update_progress("transcribe")
    with st.spinner("Transcribing audio... (This may take a while for large files)"):
        with open(audio_path, "rb") as audio_file:
            transcript_response = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file,
                response_format="verbose_json",
                timestamp_granularities=["segment"]
            )
        update_progress("transcribe", "complete")
        return transcript_response

# Function to detect if text is in English
def is_english(client, text_sample):
    update_progress("detect_language")
    with st.spinner("Detecting language..."):
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a language detection assistant. Analyze the text carefully and determine if it is primarily in English."},
                {"role": "user", "content": f"Analyze this text and determine if it is primarily in English. Answer with only 'YES' if the text is in English, or 'NO' followed by the detected language if it's not in English:\n\n{text_sample[:500]}"}
            ],
            max_tokens=20
        )
        result = response.choices[0].message.content.strip().upper()
        update_progress("detect_language", "complete")
        
        is_eng = result.startswith("YES")
        
        # Debug output
        if show_intermediate:
            if is_eng:
                st.success(f"🔍 Language detection: English confirmed")
            else:
                detected_lang = result.replace("NO", "").strip()
                st.warning(f"🔍 Language detection: Non-English detected ({detected_lang if detected_lang else 'Unknown language'})")
        
        return is_eng

# Function to translate text to English
def translate_to_english(client, text, source_language=None):
    update_progress("translate")
    with st.spinner("Translating to English..."):
        lang_prompt = f" from {source_language}" if source_language else ""
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0.1,  # Low temperature for more consistent translation
            messages=[
                {"role": "system", "content": f"""You are a professional translator specializing in technical and maintenance content.
                Translate the following text{lang_prompt} to clear, professional English while:
                1. Preserving all technical terms and procedures accurately
                2. Maintaining the chronological order of instructions
                3. Keeping any timestamps, measurements, or technical specifications exact
                4. Using clear, concise language appropriate for technical documentation
                """},
                {"role": "user", "content": f"Translate this technical content to English:\n\n{text}"}
            ]
        )
        translation = response.choices[0].message.content
        
        # Basic validation - check that we got a reasonable translation
        if len(translation.strip()) < len(text.strip()) * 0.1:  # Translation is suspiciously short
            st.warning("⚠️ Translation seems incomplete. Using original text.")
            return text
            
        update_progress("translate", "complete")
        return translation

# Function to organize content and improve structure
def organize_content(client, text, timestamps):
    update_progress("organize")
    with st.spinner("Organizing content into logical sections..."):
        # Combine text with timestamps for context
        text_with_timestamps = ""
        for i, segment in enumerate(timestamps):
            if timestamp_position == "End of steps":
                text_with_timestamps += f"{segment['text']} [{segment['start']} - {segment['end']}]\n"
            else:
                text_with_timestamps += f"[{segment['start']} - {segment['end']}] {segment['text']}\n"
            
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=temperature,
            messages=[
                {"role": "system", "content": """You are an expert in technical documentation. 
                Your task is to organize transcribed maintenance/repair instructions into logical sections.
                Maintain chronological order but group related steps together.
                Identify and correct any information that seems out of order.
                Keep all provided timestamps in your response in the same position they appear in the input."""},
                {"role": "user", "content": f"Organize the following maintenance/repair transcript, keeping timestamps:\n\n{text_with_timestamps}"}
            ]
        )
        update_progress("organize", "complete")
        return response.choices[0].message.content

# Function to generate system prompt based on fact-checking level
def get_system_prompt(mode, stage="enhance", timestamp_pos="End of steps"):
    """Generate system prompts based on fact-checking mode"""
    base_prompt = """You are an expert in technical documentation and mechanical repair."""
    
    timestamp_instruction = f"Place all timestamps at the {timestamp_pos.lower()} of each instruction step, in brackets like [00:00 - 00:00]."
    if timestamp_pos == "End of steps":
        timestamp_example = "For example, instead of \"[00:15 - 00:30] Remove the cover\", write \"Remove the cover [00:15 - 00:30]\""
    else:
        timestamp_example = "For example, write \"[00:15 - 00:30] Remove the cover\" with the timestamp at the beginning"
    
    if stage == "enhance":
        if mode == "Strict (only transcript facts)":
            return base_prompt + f"""
            Enhance maintenance instructions with ONLY information explicitly stated in the transcript.
            DO NOT add any details, steps, or explanations not directly mentioned in the source.
            {timestamp_instruction}
            {timestamp_example}
            Your goal is to organize and clarify ONLY what is explicitly stated, with no additional inferences.
            If information appears ambiguous or incomplete, state this explicitly rather than filling in gaps.
            """
        elif mode == "Balanced (minimal inference)":
            return base_prompt + f"""
            Use chain-of-thought reasoning to enhance maintenance instructions.
            You may clarify ambiguous instructions and add minimal technical context when absolutely necessary.
            When you must infer information not explicitly stated, mark it with [INFERRED: your text].
            {timestamp_instruction}
            {timestamp_example}
            Prioritize accuracy over completeness - when in doubt, exclude questionable information.
            """
        else:  # Flexible
            return base_prompt + f"""
            Use chain-of-thought reasoning to enhance maintenance instructions.
            Identify gaps in the instructions and infer missing information based on context and technical knowledge.
            Make the content more complete and understandable by adding helpful context and explanations.
            {timestamp_instruction}
            {timestamp_example}
            Use your expertise to explain WHY certain steps are important.
            While you may add helpful details, ensure they remain technically accurate and relevant.
            """
    elif stage == "manual":
        # Similar patterns for instruction manual creation...
        if mode == "Strict (only transcript facts)":
            return f"""You are an expert in creating instructional documentation.
            Transform the content into a clear instruction manual format using ONLY information provided.
            DO NOT add any details, warnings, or explanations not directly present in the input.
            Format with appropriate sections (introduction, tools, steps) but do not invent content.
            {timestamp_instruction}
            
            Use markdown formatting for headings:
            - Use # for main title
            - Use ## for section headings
            - Format steps as numbered lists
            - Make important notes bold using **text**
            """
        elif mode == "Balanced (minimal inference)":
            return f"""You are an expert in creating instructional documentation.
            Transform the content into a clear, step-by-step instruction manual format with:
            1. A title and introduction explaining what this repair/maintenance is for
            2. Tools and materials required section (based on explicit mentions only)
            3. Safety precautions section (only if explicitly mentioned)
            4. Numbered steps with clear instructions
            5. {timestamp_instruction}
            
            When you must infer information not explicitly stated, mark it with [INFERRED: your text].
            Use markdown formatting for section headings:
            - Use # for main title
            - Use ## for section headings
            - Use ### for subsection headings if needed
            - Format steps as numbered lists
            - Make important notes or warnings bold using **text**
            """
        else:  # Flexible
            return f"""You are an expert in creating instructional documentation.
            Transform the content into a comprehensive, step-by-step instruction manual format with:
            1. A title and introduction explaining what this repair/maintenance is for
            2. Tools and materials required section (infer from context)
            3. Safety precautions section (infer from context and add standard safety practices)
            4. Numbered steps with clear instructions
            5. {timestamp_instruction}
            6. Include troubleshooting tips where appropriate
            7. Create a conclusion summarizing the repair process
            
            Format everything as a professional instruction manual with markdown:
            - Use # for main title
            - Use ## for section headings
            - Use ### for subsection headings if needed
            - Format steps as numbered lists
            - Make important notes or warnings bold using **text**
            """
    
    return base_prompt

# Function to verify content against transcription (reduces hallucination)
def verify_content(client, generated_content, original_transcript, temperature=0.1):
    """Verify generated content against original transcript to reduce hallucinations"""
    update_progress("verify")
    with st.spinner("Verifying content accuracy..."):
        # Try to get a valid JSON response
        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                temperature=temperature,  # Use low temperature for verification
                response_format={"type": "json_object"},  # Explicitly request JSON format
                messages=[
                    {"role": "system", "content": """You are a verification assistant specialized in fact-checking.
                    Your task is to verify that the generated maintenance manual content is consistent with the original transcript.
                    Identify any statements in the generated content that:
                    1. Contradict the original transcript
                    2. Include specific technical details not present in or reasonably inferred from the transcript
                    3. Add steps or actions not mentioned or strongly implied in the transcript
                    
                    Do not flag general formatting improvements, clarifications, or standard technical knowledge.
                    
                    IMPORTANT: Your response MUST be a valid JSON object with exactly this structure:
                    {
                      "accurate": true/false,
                      "issues": [
                        {
                          "text": "exact problematic text",
                          "reason": "brief explanation of the issue",
                          "suggestion": "recommended fix or removal"
                        }
                      ],
                      "verified_content": "corrected content with issues fixed or flagged with [INFERRED] tags"
                    }
                    
                    If no issues are found, set "accurate" to true, provide an empty "issues" array, and return the original content.
                    DO NOT include any text before or after the JSON object.
                    """},
                    {"role": "user", "content": f"Original Transcript:\n\n{original_transcript}\n\nGenerated Content:\n\n{generated_content}"}
                ]
            )
            
            verification_result = response.choices[0].message.content
            
            # Try to parse the JSON response
            result = json.loads(verification_result)
            update_progress("verify", "complete")
            return result
        except json.JSONDecodeError:
            # If JSON parsing fails, try a second attempt with a simpler request
            if show_intermediate:
                st.warning("First verification attempt returned non-JSON. Trying again with simpler request.")
            
            try:
                # Try again with a simpler prompt focused just on accuracy
                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    temperature=0.0,  # Use zero temperature for more predictable output
                    response_format={"type": "json_object"},
                    messages=[
                        {"role": "system", "content": """Return ONLY a simple JSON object with this structure:
                        {
                          "accurate": true/false,
                          "issues": [],
                          "verified_content": "paste the original content here"
                        }
                        """},
                        {"role": "user", "content": f"Is this content accurate based on the transcript? Return JSON only.\n\nTranscript:\n{original_transcript[:1000]}...\n\nContent:\n{generated_content[:1000]}..."}
                    ]
                )
                verification_result = response.choices[0].message.content
                result = json.loads(verification_result)
                update_progress("verify", "complete")
                return result
            except Exception as e:
                # Fallback if all attempts fail
                if show_intermediate:
                    st.warning(f"Verification failed: {str(e)}. Using original content.")
                update_progress("verify", "complete")
                return {
                    "accurate": True,  # Assume accurate to avoid blocking the process
                    "issues": [],
                    "verified_content": generated_content
                }

# Function to enhance content through reasoning
def enhance_content(client, organized_text, original_transcript, temperature=0.7, fact_mode="Balanced (minimal inference)", enable_verification=True, timestamp_pos="End of steps"):
    update_progress("enhance")
    with st.spinner("Enhancing content through reasoning..."):
        # Generate the system prompt based on the fact-checking mode
        system_prompt = get_system_prompt(fact_mode, "enhance", timestamp_pos)
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=temperature,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Enhance the following maintenance instructions using technical expertise and reasoning, and place all timestamps at the {timestamp_pos.lower()} of each step:\n\n{organized_text}"}
            ]
        )
        
        enhanced_content = response.choices[0].message.content
        
        # Optional verification step
        if enable_verification:
            verification = verify_content(client, enhanced_content, original_transcript)
            
            if not verification["accurate"]:
                # Show issues if any were found
                if verification["issues"] and len(verification["issues"]) > 0 and show_intermediate:
                    st.warning("The verification process found potential inaccuracies:")
                    for issue in verification["issues"]:
                        st.info(f"**Issue:** {issue['text']}\n\n**Reason:** {issue['reason']}\n\n**Suggestion:** {issue['suggestion']}")
                
                # Use the verified content
                enhanced_content = verification["verified_content"]
        
        update_progress("enhance", "complete")
        return enhanced_content

# Function to transform into an instruction manual
def create_instruction_manual(client, enhanced_text, original_transcript, temperature=0.7, fact_mode="Balanced (minimal inference)", enable_verification=True, timestamp_pos="End of steps"):
    update_progress("create_manual")
    with st.spinner("Creating instruction manual format..."):
        # Generate the system prompt based on the fact-checking mode
        system_prompt = get_system_prompt(fact_mode, "manual", timestamp_pos)
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=temperature,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Transform the following content into a professional instruction manual with proper markdown formatting for headings, keeping timestamps at the {timestamp_pos.lower()} of steps:\n\n{enhanced_text}"}
            ]
        )
        
        instruction_manual = response.choices[0].message.content
        
        # Optional verification step
        if enable_verification:
            verification = verify_content(client, instruction_manual, original_transcript)
            
            if not verification["accurate"]:
                # Use the verified content
                instruction_manual = verification["verified_content"]
        
        update_progress("create_manual", "complete")
        return instruction_manual

# Function to extract timestamps
def extract_timestamps(text):
    pattern = r'\[(\d+:\d+(?::\d+)?) - (\d+:\d+(?::\d+)?)\]'
    matches = re.findall(pattern, text)
    return [{"start": start, "end": end} for start, end in matches]

# Function to create and format Word document with images
def create_word_document(manual_text, original_transcript, frames=None, normalized_frames=None, title="Maintenance Instruction Manual", image_size=4.0, original_language_transcript=None, include_both_transcripts=False):
    update_progress("create_document")
    with st.spinner("Creating Word document..."):
        doc = docx.Document()
        
        # Debug information about available frames
        if frames and show_intermediate:
            st.info(f"Word document has {len(frames)} images available at timestamps: {', '.join(sorted(frames.keys()))}")
        
        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add current date
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(time.strftime("%B %d, %Y"))
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        # Extract all timestamps from the text for debugging
        timestamp_pattern = r'\[(\d+:\d+(?::\d+)?) - (\d+:\d+(?::\d+)?)\]'
        all_timestamps = re.findall(timestamp_pattern, manual_text)
        if all_timestamps and show_intermediate:
            st.info(f"Found {len(all_timestamps)} timestamps in document: {', '.join([ts[0] for ts in all_timestamps])}")
        
        # Track which images we've used
        used_images = set()
        
        # Function to process a line with timestamps
        def process_line(line):
            timestamp_pattern = r'\[(\d+:\d+(?::\d+)?) - (\d+:\d+(?::\d+)?)\]'
            match = re.search(timestamp_pattern, line)
            
            if match:
                timestamp = match.group(0)
                # Always extract the start time regardless of timestamp position
                start_time = match.group(1)
                
                # Get content and timestamp position
                timestamp_index = line.find(timestamp)
                if timestamp_index > 0:  # Timestamp is not at the beginning
                    content = line[:timestamp_index].strip()
                    timestamp_at_end = True
                else:
                    content = line.replace(timestamp, "").strip()
                    timestamp_at_end = False
                
                # Create paragraph
                para = doc.add_paragraph()
                
                # Process bold text formatting in content
                bold_pattern = r'\*\*(.*?)\*\*'
                parts = re.split(bold_pattern, content)
                
                # Process inferred text formatting
                inferred_pattern = r'\[INFERRED: (.*?)\]'
                
                # Add content with proper formatting first if timestamp at end
                if timestamp_at_end:
                    for i, part in enumerate(parts):
                        # Even indices are regular text, odd indices are bold text
                        if i % 2 == 0:
                            if part:  # Only add if not empty
                                # Check for inferred text
                                inferred_matches = re.findall(inferred_pattern, part)
                                if inferred_matches:
                                    # Split by inferred text and process each piece
                                    inferred_parts = re.split(inferred_pattern, part)
                                    for j, inferred_part in enumerate(inferred_parts):
                                        if j % 2 == 0:  # Regular text
                                            if inferred_part:
                                                content_run = para.add_run(inferred_part)
                                                content_run.font.size = Pt(11)
                                        else:  # Inferred text
                                            inferred_run = para.add_run(f"[INFERRED: {inferred_part}]")
                                            inferred_run.font.size = Pt(11)
                                            inferred_run.italic = True
                                            inferred_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)  # Blue color
                                else:
                                    content_run = para.add_run(part)
                                    content_run.font.size = Pt(11)
                        else:
                            # This is text that should be bold
                            bold_run = para.add_run(part)
                            bold_run.bold = True
                            bold_run.font.size = Pt(11)
                    
                    # Add timestamp at the end
                    ts_run = para.add_run(" " + timestamp)
                    ts_run.bold = True
                    ts_run.font.size = Pt(10)
                else:
                    # Add timestamp at the beginning (legacy support)
                    ts_run = para.add_run(timestamp + " ")
                    ts_run.bold = True
                    ts_run.font.size = Pt(10)
                    
                    # Then add the content
                    for i, part in enumerate(parts):
                        # Even indices are regular text, odd indices are bold text
                        if i % 2 == 0:
                            if part:  # Only add if not empty
                                # Check for inferred text
                                inferred_matches = re.findall(inferred_pattern, part)
                                if inferred_matches:
                                    # Split by inferred text and process each piece
                                    inferred_parts = re.split(inferred_pattern, part)
                                    for j, inferred_part in enumerate(inferred_parts):
                                        if j % 2 == 0:  # Regular text
                                            if inferred_part:
                                                content_run = para.add_run(inferred_part)
                                                content_run.font.size = Pt(11)
                                        else:  # Inferred text
                                            inferred_run = para.add_run(f"[INFERRED: {inferred_part}]")
                                            inferred_run.font.size = Pt(11)
                                            inferred_run.italic = True
                                            inferred_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)  # Blue color
                                else:
                                    content_run = para.add_run(part)
                                    content_run.font.size = Pt(11)
                        else:
                            # This is text that should be bold
                            bold_run = para.add_run(part)
                            bold_run.bold = True
                            bold_run.font.size = Pt(11)
                
                # Add image if available - find exact or fuzzy match
                if frames and include_images:
                    found_image = False
                    
                    # Try direct match first
                    if start_time in frames:
                        img_para = doc.add_paragraph()
                        doc.add_picture(frames[start_time], width=Inches(image_size))
                        caption_para = doc.add_paragraph(f"Figure: Visual reference for step at {start_time}")
                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_para.style = 'Caption'
                        used_images.add(start_time)
                        found_image = True
                    
                    # If direct match fails, try normalized match
                    elif normalized_frames and not found_image:
                        normalized_start = normalize_timestamp(start_time)
                        if normalized_start in normalized_frames:
                            img_para = doc.add_paragraph()
                            doc.add_picture(normalized_frames[normalized_start], width=Inches(image_size))
                            caption_para = doc.add_paragraph(f"Figure: Visual reference for step at {start_time}")
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption_para.style = 'Caption'
                            used_images.add(normalized_start)
                            found_image = True
                    
                    # If still no match and advanced matching is enabled, try approximate matching
                    if not found_image and enable_advanced_matching:
                        # Find closest timestamp
                        closest_ts = None
                        min_diff = float('inf')
                        start_seconds = timestamp_to_seconds(start_time)
                        
                        for ts in frames.keys():
                            ts_seconds = timestamp_to_seconds(ts)
                            diff = abs(ts_seconds - start_seconds)
                            if diff < min_diff and diff <= 5:  # Within 5 seconds
                                min_diff = diff
                                closest_ts = ts
                        
                        if closest_ts:
                            img_para = doc.add_paragraph()
                            doc.add_picture(frames[closest_ts], width=Inches(image_size))
                            caption_para = doc.add_paragraph(f"Figure: Visual reference for step at {start_time} (matched with {closest_ts})")
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption_para.style = 'Caption'
                            used_images.add(closest_ts)
                            found_image = True
                
                return para
            else:
                # Regular paragraph without timestamps
                para = doc.add_paragraph()
                
                # Process bold text formatting
                bold_pattern = r'\*\*(.*?)\*\*'
                parts = re.split(bold_pattern, line)
                
                # Process inferred text formatting
                inferred_pattern = r'\[INFERRED: (.*?)\]'
                
                # Add parts with proper formatting
                for i, part in enumerate(parts):
                    # Even indices are regular text, odd indices are bold text
                    if i % 2 == 0:
                        if part:  # Only add if not empty
                            # Check for inferred text
                            inferred_matches = re.findall(inferred_pattern, part)
                            if inferred_matches:
                                # Split by inferred text and process each piece
                                inferred_parts = re.split(inferred_pattern, part)
                                for j, inferred_part in enumerate(inferred_parts):
                                    if j % 2 == 0:  # Regular text
                                        if inferred_part:
                                            content_run = para.add_run(inferred_part)
                                            content_run.font.size = Pt(11)
                                    else:  # Inferred text
                                        inferred_run = para.add_run(f"[INFERRED: {inferred_part}]")
                                        inferred_run.font.size = Pt(11)
                                        inferred_run.italic = True
                                        inferred_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)  # Blue color
                            else:
                                content_run = para.add_run(part)
                                content_run.font.size = Pt(11)
                    else:
                        # This is text that should be bold
                        bold_run = para.add_run(part)
                        bold_run.bold = True
                        bold_run.font.size = Pt(11)
                
                return para
        
        # Process the manual text
        lines = manual_text.split('\n')
        for line in lines:
            if line.strip():  # Skip empty lines
                # Check if this is a markdown heading (starts with # or ##)
                heading_match = re.match(r'^(#{1,3})\s+(.+)$', line)
                if heading_match:
                    # Get heading level and content
                    heading_level = len(heading_match.group(1))
                    heading_content = heading_match.group(2).strip()
                    
                    # Create appropriate heading style
                    if heading_level == 1:
                        # H1 heading (largest)
                        header_para = doc.add_heading(level=1)
                        header_run = header_para.add_run(heading_content)
                        header_run.font.size = Pt(16)
                    elif heading_level == 2:
                        # H2 heading
                        header_para = doc.add_heading(level=2)
                        header_run = header_para.add_run(heading_content)
                        header_run.font.size = Pt(14)
                    else:
                        # H3 heading
                        header_para = doc.add_heading(level=3)
                        header_run = header_para.add_run(heading_content)
                        header_run.font.size = Pt(12)
                # Check if this is a section header (all uppercase or ends with colon)
                elif line.isupper() or (line.strip().endswith(':') and len(line.strip()) < 50):
                    header_para = doc.add_heading(level=2)
                    header_run = header_para.add_run(line)
                    header_run.font.size = Pt(14)
                # Check if this is a numbered step (starts with a number followed by period or parenthesis)
                elif re.match(r'^\d+[\.)]\s+', line):
                    step_para = doc.add_paragraph(style='List Number')
                    step_content = re.sub(r'^\d+[\.)]\s+', '', line)
                    
                    # Check if the step content has a timestamp
                    ts_match = re.search(timestamp_pattern, step_content)
                    if ts_match:
                        # Extract timestamp for image matching
                        start_time = ts_match.group(1)
                        
                        # Process bold text in numbered steps
                        bold_pattern = r'\*\*(.*?)\*\*'
                        parts = re.split(bold_pattern, step_content)
                        
                        # Process inferred text formatting
                        inferred_pattern = r'\[INFERRED: (.*?)\]'
                        
                        # Add parts with proper formatting
                        for i, part in enumerate(parts):
                            # Even indices are regular text, odd indices are bold text
                            if i % 2 == 0:
                                if part:  # Only add if not empty
                                    # Check for inferred text
                                    inferred_matches = re.findall(inferred_pattern, part)
                                    if inferred_matches:
                                        # Split by inferred text and process each piece
                                        inferred_parts = re.split(inferred_pattern, part)
                                        for j, inferred_part in enumerate(inferred_parts):
                                            if j % 2 == 0:  # Regular text
                                                if inferred_part:
                                                    content_run = step_para.add_run(inferred_part)
                                            else:  # Inferred text
                                                inferred_run = step_para.add_run(f"[INFERRED: {inferred_part}]")
                                                inferred_run.italic = True
                                                inferred_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)  # Blue color
                                    else:
                                        content_run = step_para.add_run(part)
                            else:
                                # This is text that should be bold
                                bold_run = step_para.add_run(part)
                                bold_run.bold = True
                        
                        # Add image if available - find exact or fuzzy match
                        if frames and include_images:
                            found_image = False
                            
                            # Try direct match first
                            if start_time in frames:
                                img_para = doc.add_paragraph()
                                doc.add_picture(frames[start_time], width=Inches(image_size))
                                caption_para = doc.add_paragraph(f"Figure: Visual reference for step at {start_time}")
                                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                caption_para.style = 'Caption'
                                used_images.add(start_time)
                                found_image = True
                            
                            # If direct match fails, try normalized match
                            elif normalized_frames and not found_image:
                                normalized_start = normalize_timestamp(start_time)
                                if normalized_start in normalized_frames:
                                    img_para = doc.add_paragraph()
                                    doc.add_picture(normalized_frames[normalized_start], width=Inches(image_size))
                                    caption_para = doc.add_paragraph(f"Figure: Visual reference for step at {start_time}")
                                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    caption_para.style = 'Caption'
                                    used_images.add(normalized_start)
                                    found_image = True
                            
                            # If still no match and advanced matching is enabled, try approximate matching
                            if not found_image and enable_advanced_matching:
                                # Find closest timestamp
                                closest_ts = None
                                min_diff = float('inf')
                                start_seconds = timestamp_to_seconds(start_time)
                                
                                for ts in frames.keys():
                                    ts_seconds = timestamp_to_seconds(ts)
                                    diff = abs(ts_seconds - start_seconds)
                                    if diff < min_diff and diff <= 5:  # Within 5 seconds
                                        min_diff = diff
                                        closest_ts = ts
                                
                                if closest_ts:
                                    img_para = doc.add_paragraph()
                                    doc.add_picture(frames[closest_ts], width=Inches(image_size))
                                    caption_para = doc.add_paragraph(f"Figure: Visual reference for step at {start_time} (matched with {closest_ts})")
                                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    caption_para.style = 'Caption'
                                    used_images.add(closest_ts)
                                    found_image = True
                    else:
                        # Regular step without timestamp
                        # Process bold text in numbered steps
                        bold_pattern = r'\*\*(.*?)\*\*'
                        parts = re.split(bold_pattern, step_content)
                        
                        # Add parts with proper formatting
                        for i, part in enumerate(parts):
                            # Even indices are regular text, odd indices are bold text
                            if i % 2 == 0:
                                if part:  # Only add if not empty
                                    step_para.add_run(part)
                            else:
                                # This is text that should be bold
                                bold_run = step_para.add_run(part)
                                bold_run.bold = True
                else:
                    process_line(line)
        
        # Add original transcript as reference
        doc.add_page_break()
        
        # If we have both original and translated transcripts, add both
        if include_both_transcripts and original_language_transcript and original_language_transcript != original_transcript:
            # Add original language transcript first
            original_para = doc.add_heading(level=1)
            original_run = original_para.add_run("ORIGINAL LANGUAGE TRANSCRIPT")
            original_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            for line in original_language_transcript.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            
            doc.add_paragraph()  # Add some space
            doc.add_paragraph()
            
            # Add translated transcript
            translated_para = doc.add_heading(level=1)
            translated_run = translated_para.add_run("ENGLISH TRANSLATION")
            translated_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            for line in original_transcript.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
        else:
            # Add only the working transcript (original or translated)
            reference_para = doc.add_heading(level=1)
            reference_run = reference_para.add_run("TRANSCRIPT REFERENCE")
            reference_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            for line in original_transcript.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
                
        # Convert to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        # Show image usage summary
        if frames and include_images and show_intermediate:
            st.info(f"Used {len(used_images)} of {len(frames)} available images in the document")
            if len(used_images) < len(frames):
                # Show only the first 10 unused timestamps to avoid overwhelming the UI
                unused_images = set(frames.keys()) - used_images
                display_unused = list(sorted(unused_images))[:10]
                st.warning(f"First 10 unused timestamps: {', '.join(display_unused)}")
                
                # Show which timestamps in the document didn't get images
                document_ts = [ts[0] for ts in all_timestamps]
                unmatched_ts = []
                for ts in document_ts:
                    normalized_ts = normalize_timestamp(ts)
                    if ts not in used_images and normalized_ts not in used_images:
                        unmatched_ts.append(ts)
                
                if unmatched_ts:
                    st.error(f"Timestamps in document that couldn't be matched to images: {', '.join(unmatched_ts[:10])}")
        
        update_progress("create_document", "complete")
        return docx_bytes, used_images

# Function to create a fallback document with forced image insertion
def create_fallback_document(manual_text, original_transcript, frames, title="Maintenance Manual (with images)", image_size=4.0, original_language_transcript=None, include_both_transcripts=False):
    """Create a fallback document that forces image insertion at regular intervals"""
    with st.spinner("Creating fallback document with forced image insertion..."):
        doc = docx.Document()
        
        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title + " (WITH IMAGES)")
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add current date
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(time.strftime("%B %d, %Y"))
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        # Get all frame timestamps and sort them
        frame_timestamps = sorted(list(frames.keys()), key=lambda ts: timestamp_to_seconds(ts))
        
        # Process the manual text - use similar logic to your existing function
        # but insert images at regular intervals throughout the document
        
        # Split content into sections (we'll assume anything with a heading is a new section)
        sections = []
        current_section = []
        
        lines = manual_text.split('\n')
        for line in lines:
            if re.match(r'^#{1,3}\s+.+$', line) or not current_section:  # New section or first line
                if current_section:  # Don't append empty first section
                    sections.append(current_section)
                current_section = [line]
            else:
                current_section.append(line)
        
        # Add the last section if it exists
        if current_section:
            sections.append(current_section)
        
        # Calculate how many images to use per section
        total_frames = len(frame_timestamps)
        total_sections = len(sections)
        
        images_per_section = max(1, min(3, total_frames // total_sections))
        
        # Process each section and add images
        current_frame_index = 0
        
        for i, section in enumerate(sections):
            # Process section content
            for line in section:
                # Process the line (simplified - you should use your full processing logic here)
                if re.match(r'^#{1,3}\s+(.+)$', line):
                    # Heading
                    heading_match = re.match(r'^(#{1,3})\s+(.+)$', line)
                    heading_level = len(heading_match.group(1))
                    heading_content = heading_match.group(2).strip()
                    
                    header_para = doc.add_heading(level=heading_level)
                    header_run = header_para.add_run(heading_content)
                    header_run.font.size = Pt(18 - (heading_level * 2))  # Adjust size based on level
                else:
                    # Regular content
                    para = doc.add_paragraph(line)
            
            # Add images after this section if we have frames left
            images_to_add = min(images_per_section, total_frames - current_frame_index)
            
            for j in range(images_to_add):
                if current_frame_index < total_frames:
                    ts = frame_timestamps[current_frame_index]
                    img_path = frames[ts]
                    
                    # Add the image
                    doc.add_paragraph()  # Spacer
                    doc.add_picture(img_path, width=Inches(image_size))
                    
                    # Add a caption
                    caption_para = doc.add_paragraph(f"Figure: Image at timestamp {ts}")
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_para.style = 'Caption'
                    
                    current_frame_index += 1
        
        # Add original transcript as reference
        doc.add_page_break()
        
        # If we have both original and translated transcripts, add both
        if include_both_transcripts and original_language_transcript and original_language_transcript != original_transcript:
            # Add original language transcript first
            original_para = doc.add_heading(level=1)
            original_run = original_para.add_run("ORIGINAL LANGUAGE TRANSCRIPT")
            original_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            for line in original_language_transcript.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            
            doc.add_paragraph()  # Add some space
            doc.add_paragraph()
            
            # Add translated transcript
            translated_para = doc.add_heading(level=1)
            translated_run = translated_para.add_run("ENGLISH TRANSLATION")
            translated_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            for line in original_transcript.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
        else:
            # Add only the working transcript (original or translated)
            reference_para = doc.add_heading(level=1)
            reference_run = reference_para.add_run("TRANSCRIPT REFERENCE")
            reference_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            for line in original_transcript.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
        
        # Convert to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes

# Display processing status with a nice UI
def create_processing_status():
    stages = [
        {"name": "extract_audio", "label": "Audio Extraction", "show_for": ["mp4"]},
        {"name": "extract_frames", "label": "Image Extraction", "show_for": ["mp4"]},
        {"name": "transcribe", "label": "Transcription", "show_for": ["mp3", "mp4"]},
        {"name": "detect_language", "label": "Language Detection", "show_for": ["mp3", "mp4"]},
        {"name": "translate", "label": "Translation", "show_for": ["mp3", "mp4"]},
        {"name": "organize", "label": "Content Organization", "show_for": ["mp3", "mp4"]},
        {"name": "enhance", "label": "Content Enhancement", "show_for": ["mp3", "mp4"]},
        {"name": "verify", "label": "Verification", "show_for": ["mp3", "mp4"]},
        {"name": "create_manual", "label": "Manual Creation", "show_for": ["mp3", "mp4"]},
        {"name": "create_document", "label": "Document Generation", "show_for": ["mp3", "mp4"]}
    ]
    
    # Determine file type
    file_type = None
    if uploaded_file:
        if uploaded_file.name.endswith('.mp3'):
            file_type = "mp3"
        elif uploaded_file.name.endswith('.mp4'):
            file_type = "mp4"
    
    # Only display relevant stages based on file type
    if file_type:
        st.markdown("### Processing Status")
        cols = st.columns(2)
        
        for i, stage in enumerate(stages):
            if file_type in stage["show_for"]:
                col = cols[i % 2]
                with col:
                    status = st.session_state.get(f"{stage['name']}_status", "")
                    current = st.session_state.processing_stage == stage["name"]
                    
                    if status == "complete":
                        st.success(f"✅ {stage['label']}")
                    elif current:
                        st.info(f"🔄 {stage['label']} (in progress)")
                    elif status:
                        st.warning(f"⚠️ {stage['label']} ({status})")
                    else:
                        st.text(f"⏳ {stage['label']} (pending)")

# Main application flow
if start_processing:
    # Reset progress tracking
    for stage in ["extract_audio", "extract_frames", "transcribe", "detect_language", 
                 "translate", "organize", "enhance", "verify", "create_manual", "create_document"]:
        if f"{stage}_status" in st.session_state:
            del st.session_state[f"{stage}_status"]
    
    # Display processing status
    create_processing_status()
    
    # Save the uploaded file to a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        temp_file_path = tmp_file.name
    
    # Initialize frames dictionary for storing images
    frames = None
    normalized_frames = None
        
    try:
        # Extract audio if file is MP4
        if uploaded_file.name.endswith('.mp4'):
            audio_path = extract_audio(temp_file_path)
        else:
            audio_path = temp_file_path
            
        # Transcribe the audio
        transcript_response = transcribe_audio(client, audio_path)
        
        # Combine all segments to get full text
        full_transcript = ""
        timestamps = []
        
        if hasattr(transcript_response, 'segments'):
            for segment in transcript_response.segments:
                full_transcript += segment.text + " "
                timestamps.append({
                    "start": time.strftime('%M:%S', time.gmtime(segment.start)),
                    "end": time.strftime('%M:%S', time.gmtime(segment.end)),
                    "text": segment.text
                })
        else:
            # Fallback if segment data isn't available
            full_transcript = transcript_response.text
            
        # Display the transcript
        if show_intermediate:
            st.subheader("Raw Transcript")
            st.write(full_transcript)
        
        # Extract frames from video if it's an MP4
        if uploaded_file.name.endswith('.mp4') and include_images:
            frames, normalized_frames = extract_frames_at_timestamps(temp_file_path, timestamps)
            
            # Run diagnostics if enabled
            if show_intermediate and enable_advanced_matching:
                debug_timestamp_formats(timestamps, frames)
                visualize_timestamp_matches(timestamps, frames)
            
            # Display extracted images
            if show_intermediate:
                display_extracted_images(frames)
        
        # Check if transcript is in English and translate if necessary
        english_transcript = full_transcript
        needs_translation = not is_english(client, full_transcript[:1000])
        
        if show_intermediate:
            st.subheader("Language Detection Results")
            if needs_translation:
                st.warning(f"🌐 Detected non-English content. Translation required.")
                st.info("Original transcript preview (first 200 chars):")
                st.code(full_transcript[:200] + "...")
            else:
                st.success("✅ Content is already in English.")
        
        if needs_translation:
            if show_intermediate:
                st.info("🔄 Translating to English...")
            english_transcript = translate_to_english(client, full_transcript)
            if show_intermediate:
                st.subheader("📝 English Translation")
                st.success("✅ Translation completed successfully!")
                st.info("Translated transcript preview (first 200 chars):")
                st.code(english_transcript[:200] + "...")
                
                # Show comparison
                with st.expander("📊 View Full Translation Comparison"):
                    col1, col2 = st.columns(2)
                    with col1:
                        st.write("**Original:**")
                        st.write(full_transcript[:1000] + ("..." if len(full_transcript) > 1000 else ""))
                    with col2:
                        st.write("**Translation:**")
                        st.write(english_transcript[:1000] + ("..." if len(english_transcript) > 1000 else ""))
        
        # Organize content
        organized_text = organize_content(client, english_transcript, timestamps)
        
        if show_intermediate:
            st.subheader("📋 Content Organization")
            st.success("✅ Content organized successfully!")
            if needs_translation:
                st.info("ℹ️ Using translated English content for organization.")
            with st.expander("View Organized Content"):
                st.write(organized_text)
        
        # Enhance content
        enhanced_text = enhance_content(
            client, 
            organized_text, 
            english_transcript,
            temperature=temperature,
            fact_mode=fact_checking,
            enable_verification=enable_verification,
            timestamp_pos=timestamp_position
        )
        
        if show_intermediate:
            st.subheader("✨ Content Enhancement")
            st.success("✅ Content enhanced successfully!")
            if needs_translation:
                st.info("ℹ️ Enhanced content is based on translated English transcript.")
            with st.expander("View Enhanced Content"):
                # Highlight inferred content if that option is enabled
                if highlight_inferences and "[INFERRED:" in enhanced_text:
                    highlighted_text = enhanced_text.replace("[INFERRED:", "**[INFERRED:**").replace("]", "**]**")
                    st.markdown(highlighted_text)
                else:
                    st.write(enhanced_text)
        
        # Create instruction manual
        instruction_manual = create_instruction_manual(
            client, 
            enhanced_text, 
            english_transcript,
            temperature=temperature,
            fact_mode=fact_checking,
            enable_verification=enable_verification,
            timestamp_pos=timestamp_position
        )
        
        st.subheader("📖 Final Instruction Manual")
        if needs_translation:
            st.info("✅ Manual created in English from translated content")
        else:
            st.info("✅ Manual created in English from original content")
            
        # Highlight inferred content if that option is enabled
        if highlight_inferences and "[INFERRED:" in instruction_manual:
            highlighted_manual = instruction_manual.replace("[INFERRED:", "**[INFERRED:**").replace("]", "**]**")
            st.markdown(highlighted_manual)
        else:
            st.write(instruction_manual)
        
        # Set default title if not provided
        if not manual_title:
            manual_title = f"Maintenance Manual for {uploaded_file.name.split('.')[0]}"
        
        # Create Word document
        word_doc, used_images = create_word_document(
            instruction_manual, 
            english_transcript, 
            frames=frames if include_images else None, 
            normalized_frames=normalized_frames if include_images else None, 
            title=manual_title,
            image_size=image_size,
            original_language_transcript=full_transcript,
            include_both_transcripts=include_original_transcript
        )
        
        # Check if no images were used
        if frames and include_images and len(used_images) == 0:
            st.warning("⚠️ No images were matched in the document. Trying emergency fallback mode...")
            
            # Emergency fallback - force image insertion for each step
            fallback_doc = create_fallback_document(
                instruction_manual,
                english_transcript,
                frames,
                title=manual_title,
                image_size=image_size,
                original_language_transcript=full_transcript,
                include_both_transcripts=include_original_transcript
            )
            
            st.success("✅ Created fallback document with forced image matching.")
            
            # Offer both documents
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    label="📥 Download Original Document",
                    data=word_doc,
                    file_name=f"{manual_title.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            
            with col2:
                st.download_button(
                    label="📥 Download Fallback Document (with images)",
                    data=fallback_doc,
                    file_name=f"{manual_title.replace(' ', '_')}_with_images.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    type="primary"
                )
        else:
            # Standard download button
            st.download_button(
                label="📥 Download Word Document",
                data=word_doc,
                file_name=f"{manual_title.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary"
            )
        
        # Success message
        st.success("✅ Processing complete! Download your maintenance manual above.")
        
    except Exception as e:
        st.error(f"An error occurred: {str(e)}")
        update_progress("error")
    
    finally:
        # Clean up temporary files
        if os.path.exists(temp_file_path):
            os.unlink(temp_file_path)
        if 'audio_path' in locals() and audio_path != temp_file_path and os.path.exists(audio_path):
            os.unlink(audio_path)
        # Clean up image frames directory if it exists
        if frames:
            # Get the temp directory path
            temp_dir_path = os.path.dirname(next(iter(frames.values())))
            # Use shutil.rmtree to recursively remove directory and all contents
            if os.path.exists(temp_dir_path):
                try:
                    shutil.rmtree(temp_dir_path)
                except Exception as e:
                    st.warning(f"Could not remove temporary directory: {str(e)}")

elif st.session_state.processing_stage is not None:
    # Continue displaying the processing status during ongoing processing
    create_processing_status()

# If no processing has started yet, show the introduction
else:
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown("""
        ## Convert maintenance videos to instruction manuals
        
        This application helps maintenance teams transform video/audio recordings into 
        structured documentation with timestamps and images.
        
        ### How it works:
        1. Upload a maintenance or repair video/audio file
        2. AI transcribes the content and analyzes the instructions
        3. For videos, images are extracted at key timestamps
        4. A structured instruction manual is generated with:
           - Clear step-by-step instructions
           - Timestamps for each step
           - Images showing the process (for videos)
           - Professional formatting
           
        ### Getting started:
        1. Enter your OpenAI API key in the sidebar
        2. Upload an MP3 or MP4 file
        3. Adjust settings as needed
        4. Click "Start Processing"
        """)
    
    with col2:
        st.image("logo/Maintenance Manual Generator.png", use_container_width=True)

# Display footer
st.markdown("---")
st.caption("© 2025 Maintenance Manual Generator | Uses OpenAI API for transcription and processing")
